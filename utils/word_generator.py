from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from math import ceil

def next_multiple_of_5(n):
    return int(ceil(n / 5.0)) * 5

def add_pack_section(doc, index, customer, part, desc, qty, po, mfg, qty_line):
    doc.add_paragraph(f"{index}. {customer} {part} Pack Label", style="Heading2")
    doc.add_paragraph(f"Part Number: {part}")
    doc.add_paragraph(f"Description: {desc}")
    doc.add_paragraph(f"Quantity: {qty} PCS")
    doc.add_paragraph(f"PO Number: {po}")
    doc.add_paragraph("Bar Code:")
    doc.add_paragraph(f"Mfg. {mfg}")
    doc.add_paragraph(f"Quantity: {qty_line}")
    doc.add_page_break()

def generate_word_doc(df, path, mfg, po, customer):
    doc = Document()
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(11)

    for _, row in df.iterrows():
        no = str(row['no'])
        part = str(row['part_number'])
        desc = str(row['description'])
        units_per_pack = int(row.get('units_per_pack', 0))
        req_packs = int(row.get('req_packs', 0))
        odd_packs = str(row.get('odd_packs', "")).strip()
        odd_units = int(row.get('odd_units', 0))

        index = f"{no}.1"
        add_pack_section(
            doc, index, customer, part, desc, units_per_pack, po, mfg,
            next_multiple_of_5(req_packs)
        )

        if odd_packs and odd_packs != "0":
            index2 = f"{no}.2"
            add_pack_section(
                doc, index2, customer, part, desc, odd_units, po, mfg, 3
            )

    doc.save(path)
