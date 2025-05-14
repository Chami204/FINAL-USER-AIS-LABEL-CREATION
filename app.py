import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Inches
import barcode
from barcode.writer import ImageWriter
import os
import shutil
from zipfile import ZipFile

# Set page config
st.set_page_config(page_title="Packaging Label Generator", layout="centered", page_icon="📦")

st.markdown(
    "<h1 style='color:#023E8A;text-align:center;'>📦 Packaging Label Generator</h1>",
    unsafe_allow_html=True
)

# Step 1: Table input
st.subheader("Step 1: Enter Packaging Details")
num_rows = st.number_input("How many rows do you want to input?", min_value=1, max_value=100, value=1)

table_data = {
    "No.": ["" for _ in range(num_rows)],
    "Part Number": ["" for _ in range(num_rows)],
    "Description": ["" for _ in range(num_rows)],
    "Quantity per pack": ["" for _ in range(num_rows)],
    "Number of packs": ["" for _ in range(num_rows)]
}

df = pd.DataFrame(table_data)
edited_df = st.data_editor(df, num_rows="dynamic", use_container_width=True)

# Step 2: Additional inputs
st.subheader("Step 2: Enter Additional Information")
customer = st.text_input("Customer Name")
po_number = st.text_input("PO Number")
mfg = st.text_input("Mfg. (e.g. 03-2025)")

if st.button("Generate Packaging Labels"):

    # Validate
    if edited_df.isnull().values.any() or not customer or not po_number or not mfg:
        st.error("Please fill in all the fields before generating labels.")
    else:
        # Prepare output directories
        output_dir = "output"
        barcode_dir = os.path.join(output_dir, "barcodes")
        os.makedirs(barcode_dir, exist_ok=True)

        # Create Word document
        doc = Document()

        for index, row in edited_df.iterrows():
            no = str(row['No.']).strip()
            part_number = str(row['Part Number']).strip()
            description = str(row['Description']).strip()
            qty_pack = str(row['Quantity per pack']).strip()
            num_packs = str(row['Number of packs']).strip()

            # Add to Word
            doc.add_paragraph(f"{no}. {customer} {part_number} Pack Label")
            doc.add_paragraph(f"Part Number: {part_number}")
            doc.add_paragraph(f"Description: {description}")
            doc.add_paragraph(f"Quantity: {qty_pack} PCS")
            doc.add_paragraph(f"PO Number: {po_number}")
            doc.add_paragraph(f"Bar Code:")
            doc.add_paragraph(f"Mfg. {mfg}")
            doc.add_paragraph(f"Quantity: {num_packs}")
            doc.add_page_break()

            # Generate barcode
            barcode_text = part_number
            barcode_filename = f"{no}.{customer}.{part_number}.PackLabel".replace(" ", "_")
            barcode_path = os.path.join(barcode_dir, barcode_filename)
            code128 = barcode.get("code128", barcode_text, writer=ImageWriter())
            code128.save(barcode_path)

        # Save Word doc
        doc_path = os.path.join(output_dir, "Packaging_Labels.docx")
        doc.save(doc_path)

        # Create zip file
        zip_path = os.path.join(output_dir, "Packaging_Labels.zip")
        with ZipFile(zip_path, 'w') as zipf:
            zipf.write(doc_path, arcname="Packaging_Labels.docx")
            for file in os.listdir(barcode_dir):
                zipf.write(os.path.join(barcode_dir, file), arcname=f"barcodes/{file}")

        with open(zip_path, "rb") as f:
            st.success("✅ Labels and Barcodes Generated!")
            st.download_button("📥 Download All Files (ZIP)", f, file_name="Packaging_Labels.zip")
